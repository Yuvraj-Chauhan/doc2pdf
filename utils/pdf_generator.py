from pikepdf import Pdf
import aspose.words as aw
from pathlib import Path

def convert_to_pdf(docx_path):
    pdf_path = Path(docx_path).with_suffix(".pdf")
    try:
        doc = aw.Document(str(docx_path))
        doc.save(str(pdf_path), aw.SaveFormat.PDF)
        return pdf_path
    except Exception as e:
        raise RuntimeError(f"Error converting DOCX to PDF: {e}")

'''
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from pathlib import Path
def convert_to_pdf(docx_path):
    pdf_path = Path(docx_path).with_suffix(".pdf")
    doc = Document(docx_path)

    # Create a ReportLab PDF object
    c = canvas.Canvas(str(pdf_path), pagesize=letter)

    width, height = letter  # Get the dimensions of the letter page
    margin = 40  # Left margin
    top = height - margin  # Start at the top of the page
    line_height = 14  # Line height for text

    # Set the font and size for the PDF
    c.setFont("Helvetica", 10)

    # Iterate through the paragraphs in the DOCX file
    for para in doc.paragraphs:
        text = para.text
        # Add text to the PDF, adjusting for new lines
        c.drawString(margin, top, text)
        top -= line_height

        # If the text runs off the page, create a new page
        if top < margin:
            c.showPage()
            c.setFont("Helvetica", 10)
            top = height - margin

    # Save the PDF file
    c.save()

    return pdf_path
'''
# from fpdf import FPDF
# from pathlib import Path
# def convert_to_pdf(docx_path):
#     pdf_path = Path(docx_path).with_suffix(".pdf")
#     pdf = FPDF()
#     pdf.set_auto_page_break(auto=True, margin=15)
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)

#     try:
#         # Load DOCX file
#         doc = Document(docx_path)

#         # Loop through paragraphs
#         for paragraph in doc.paragraphs:
#             # Encode text to handle special characters and decode them back to Unicode
#             paragraph_text = paragraph.text.encode("latin-1", errors="replace").decode("latin-1")
#             # Add paragraph text to PDF, wrapping long text
#             pdf.multi_cell(0, 10, paragraph_text)

#         # Output the PDF
#         pdf.output(str(pdf_path))
#         return pdf_path

#     except Exception as e:
#         raise RuntimeError(f"Error converting DOCX to PDF: {e}")
    
def add_password_protection(pdf_path, password):
    protected_path = pdf_path.with_name(f"{pdf_path.stem}_protected.pdf")
    with Pdf.open(pdf_path) as pdf:
        pdf.save(protected_path, encryption=Pdf.Encryption(owner=password, user=password))
    return protected_path
